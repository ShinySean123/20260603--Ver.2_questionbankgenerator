import streamlit as st
import pandas as pd
import re
import math
import io
import json
import time      
import random    
import urllib.request
import urllib.parse
import requests  
import base64
import fitz  # PyMuPDF 高畫質影像渲染引擎
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter

# Word 處理相關
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== 0. 全域常規樣式與字元定義 ====================
# 使用 chr(96) * 3 動態拼接三個反引號，100% 避免 Git 與 Streamlit Cloud 語法截斷 Bug
TRIPLE_BACKTICK = chr(96) * 3
BT_JSON = TRIPLE_BACKTICK + "json"
BT_ONLY = TRIPLE_BACKTICK

# GitHub 連動參數 (設為全域常數，利於跨模組大備份呼叫)
GITHUB_USER = "ShinySean123"
GITHUB_REPO = "20260603--Ver.2_questionbankgenerator"
GITHUB_FOLDER_HIST = "history_db"          
GITHUB_FOLDER_PDF = "current_materials"    

def sanitize_f(name): 
    """全域共用的檔名非法字元過濾器"""
    return re.sub(r'[\\/:*?"<>|]', '_', str(name))

# 🌟 [全自動雲端儲存]：備份推送至 GitHub 歷史庫核心引擎
def upload_excel_to_github(file_bytes, file_name, github_token):
    if not github_token:
        return False, "未設定 GitHub Token"
    
    encoded_user = urllib.parse.quote(GITHUB_USER)
    encoded_repo = urllib.parse.quote(GITHUB_REPO)
    sanitized_name = sanitize_f(file_name)
    encoded_path = urllib.parse.quote(f"{GITHUB_FOLDER_HIST}/{sanitized_name}")
    
    url = f"https://api.github.com/repos/{encoded_user}/{encoded_repo}/contents/{encoded_path}"
    headers = {
        "Authorization": f"token {github_token}",
        "Accept": "application/vnd.github.v3+json",
        "User-Agent": "Streamlit-App"
    }
    
    # 1. 嘗試獲取現有檔案的 SHA 雜湊 (以利覆蓋更新，避免 GitHub API 衝突)
    sha = None
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code == 200:
            sha = resp.json().get("sha")
    except Exception:
        pass
        
    # 2. 進行 Base64 編碼
    b64_content = base64.b64encode(file_bytes).decode("utf-8")
    payload = {
        "message": f"🤖 AI 自動同步備份題庫: {sanitized_name}",
        "content": b64_content,
        "branch": "main"
    }
    if sha:
        payload["sha"] = sha
        
    # 3. 發送 PUT 請求存入庫
    try:
        resp = requests.put(url, headers=headers, json=payload, timeout=15)
        if resp.status_code in [200, 201]:
            return True, f"成功！檔案已即時同步備份至 GitHub `history_db/{sanitized_name}`"
        elif resp.status_code == 401:
            return False, "GitHub 認證失敗 (401)。後台 Secrets 中的 Token 可能已被註銷，請更換新 Token。"
        else:
            return False, f"GitHub 拒絕寫入 ({resp.status_code}): {resp.text}"
    except Exception as e:
        return False, f"連線異常: {e}"

EXCEL_COL_WIDTHS = {
    'A': 8,   # 題號
    'B': 45,  # 題目內容
    'C': 30,  # 選項A
    'D': 30,  # 選項B
    'E': 30,  # 選項C
    'F': 30,  # 選項D
    'G': 30,  # 選項E
    'H': 15,  # 正確答案
    'I': 60,  # 針對各選項之詳解
    'J': 40   # 出處
}
EXCEL_BORDER = Border(
    top=Side(style='thin'), 
    bottom=Side(style='thin'), 
    left=Side(style='thin'), 
    right=Side(style='thin')
)

if "last_api_call_time" not in st.session_state:
    st.session_state["last_api_call_time"] = 0.0

# 網頁配置
st.set_page_config(page_title="AI 醫學共筆題庫工作站", page_icon="🧠", layout="centered")

st.title("🧠 AI 醫學共筆題庫三模工作站")
st.markdown("共筆組長專屬完全體：整合【講義智慧出題】、【純題配詳解】與【現成題庫轉 Word】三大核心功能！")

# ==================== 1. 🔑 共享 API 金鑰與 GitHub 後台備份設定 ====================
env_key = ""
github_token = ""

try:
    if "GEMINI_API_KEY" in st.secrets: 
        env_key = st.secrets["GEMINI_API_KEY"]
    # 🌟 徹底移入後台：優先且只從 Streamlit Secrets 安全通道讀取，絕不暴露於前端
    if "GITHUB_TOKEN" in st.secrets:
        github_token = st.secrets["GITHUB_TOKEN"].strip()
except Exception: 
    pass

with st.sidebar:
    st.header("🔑 API 金鑰配置")
    user_live_key = st.text_input("請輸入 Gemini API Key：", value=env_key if env_key else "", type="password")
    api_key = user_live_key.strip() if user_live_key else env_key
    
    st.markdown("---")
    st.header("☁️ GitHub 雲端備份狀態")
    # 🌟 在側邊欄僅展示同步狀態，不提供輸入框，保持介面清爽安全
    if github_token:
        st.success("🟢 雲端同步備份已啟動 (已偵測到後台 GITHUB_TOKEN)")
    else:
        st.info("ℹ️ 雲端同步備份未啟動 (未偵測到後台 GITHUB_TOKEN)")
        
    st.markdown("---")
    st.caption("💡 提示：『功能 A'』與『模組 C』為本地純文字引擎與排版引擎，完全不需要輸入 Gemini Key 即可完美運作！若需自動同步備份，請至 Streamlit 後台 Settings ➡️ Secrets 中配置 `GITHUB_TOKEN`。")

# 導覽器放置在金鑰檢查前，確保介面正常渲染
main_mode = st.radio(
    "🎯 請選擇您目前想要使用的共筆功能模組：",
    [
        "📚 模組 A：講義圖文智慧出題", 
        "📝 模組 B：現成題目自動配詳解", 
        "📄 模組 C：既有題庫 Excel/JSON ➡️ 轉 Word/Excel"
    ],
    index=0,
    horizontal=True
)

st.markdown("---")

if not api_key and main_mode in ["📚 模組 A：講義圖文智慧出題", "📝 模組 B：現成題目自動配詳解"]:
    st.warning("⚠️ 請先在左側邊欄填入您在 Google AI Studio 申請的 `AIzaSy` 金鑰以解鎖系統。")
    st.stop()

# ==================== 🌟 共享的終極單發 HTTP 直連函數 ====================
def generate_content_via_http_with_retry(contents_list, api_key, max_retries=4):
    models_pool = ["gemini-2.5-flash", "gemini-3-flash", "gemini-3.5-flash"]
    parts = []
    for item in contents_list:
        parts.append({"text": str(item)})
    payload = {"contents": [{"parts": parts}]}
    headers = {"Content-Type": "application/json"}

    for model_name in models_pool:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=90)
            if resp.status_code == 200:
                st.session_state["last_api_call_time"] = time.time()
                return resp.json()['candidates'][0]['content']['parts'][0]['text']
            elif resp.status_code in [429, 503]:
                continue
        except Exception:
            continue

    st.warning("🔄 雲端排隊人數較多，正在為您進行跨世代核心調節調度...")
    for attempt in range(max_retries):
        current_time = time.time()
        elapsed = current_time - st.session_state["last_api_call_time"]
        REQUIRED_GAP = 12.5
        if elapsed < REQUIRED_GAP:
            time.sleep(REQUIRED_GAP - elapsed)
            
        for model_name in models_pool:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
            try:
                resp = requests.post(url, headers=headers, json=payload, timeout=90)
                if resp.status_code == 200:
                    st.session_state["last_api_call_time"] = time.time()
                    return resp.json()['candidates'][0]['content']['parts'][0]['text']
            except Exception:
                pass
        time.sleep((2 ** attempt) + random.uniform(0, 1))
    raise Exception("Google 伺服器目前全線大塞車，請稍候幾分鐘再試。")

# ==============================================================================
# 🌟 模組 A：全自動講義圖文出題系統
# ==============================================================================
if "模組 A" in main_mode:
    st.subheader("📚 模式 A：課程講義出題工作站")
    sub_function_mode = st.radio(
        "🔥 請選擇您的智慧出題執行方式：",
        ["✅ 功能 A：直接點擊 ➡️ 網頁自動呼叫雲端 AI 一鍵吐出檔案", "🚀 功能 A'：客製化生成專屬 PROMPT ➡️ 自己複製拿去餵給任意大模型 AI"],
        index=0,
        horizontal=True
    )
    st.markdown("---")

    encoded_user = urllib.parse.quote(GITHUB_USER)
    encoded_repo = urllib.parse.quote(GITHUB_REPO)

    github_api_hist_url = f"https://api.github.com/repos/{encoded_user}/{encoded_repo}/contents/{urllib.parse.quote(GITHUB_FOLDER_HIST)}"
    file_options = ["❌ 不使用歷史資料（全新出題）"]
    all_excel_files = [] 

    try:
        req = urllib.request.Request(github_api_hist_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response:
            api_data = json.loads(response.read().decode())
        for item in api_data:
            if item['type'] == 'file' and item['name'].endswith('.xlsx'): all_excel_files.append(item['name'])
    except Exception:
        try:
            html_url = f"https://github.com/{encoded_user}/{encoded_repo}/tree/main/{urllib.parse.quote(GITHUB_FOLDER_HIST)}"
            req = urllib.request.Request(html_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as resp: html_text = resp.read().decode('utf-8')
            all_excel_files = list(set(re.findall(r'title="([^"]+\.xlsx)"', html_text)))
        except: pass

    if all_excel_files:
        file_options.append("💥 比對資料夾內【所有檔案】（全面防重複）")
        for f in all_excel_files: file_options.append(f)

    cloud_pdf_files = []
    github_api_pdf_url = f"https://api.github.com/repos/{encoded_user}/{encoded_repo}/contents/{urllib.parse.quote(GITHUB_FOLDER_PDF)}"

    try:
        req = urllib.request.Request(github_api_pdf_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response:
            pdf_api_data = json.loads(response.read().decode())
        for item in pdf_api_data:
            if item['type'] == 'file' and item['name'].lower().endswith('.pdf'): cloud_pdf_files.append(item['name'])
    except Exception:
        try:
            html_url = f"https://github.com/{encoded_user}/{encoded_repo}/tree/main/{urllib.parse.quote(GITHUB_FOLDER_PDF)}"
            req = urllib.request.Request(html_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as resp: 
                html_text = resp.read().decode('utf-8')
            found_pdfs = re.findall(r'title="([^"]+\.[pP][dD][fF])"', html_text)
            cloud_pdf_files = list(set([urllib.parse.unquote(p) for p in found_pdfs]))
        except Exception: pass

    history_titles = []
    with st.sidebar:
        st.header("⚙️ 雲端書櫃與歷史庫狀態")
        selected_mode = st.selectbox("請選擇歷史題庫防重複模式：", file_options)
        st.markdown("---")
        if cloud_pdf_files:
            st.success(f"🟢 偵測到雲端書櫃有 {len(cloud_pdf_files)} 份 PDF")
            selected_cloud_pdfs = st.multiselect("請勾選本次想連動出題的雲端講義：", cloud_pdf_files)
        else:
            st.info("ℹ️ 目前雲端書櫃內未偵測到 PDF。")
            selected_cloud_pdfs = []

    def fetch_excel_titles(file_name):
        encoded_name = urllib.parse.quote(file_name)
        raw_url = f"https://raw.githubusercontent.com/{encoded_user}/{encoded_repo}/main/{GITHUB_FOLDER_HIST}/{encoded_name}"
        try:
            req = urllib.request.Request(raw_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as resp: 
                df = pd.read_excel(io.BytesIO(resp.read()))
            q_col = next((c for c in df.columns if any(k in str(c) for k in ["題目", "Question"])), None)
            if q_col: return df[q_col].dropna().astype(str).tolist()
        except: pass
        return []

    if "【所有檔案】" in selected_mode:
        for f in all_excel_files: history_titles.extend(fetch_excel_titles(f))
    elif selected_mode != "❌ 不使用歷史資料（全新出題）":
        history_titles = fetch_excel_titles(selected_mode)

    def fetch_cloud_pdf_bytes(file_name):
        encoded_name = urllib.parse.quote(file_name)
        raw_url = f"https://raw.githubusercontent.com/{encoded_user}/{encoded_repo}/main/{GITHUB_FOLDER_PDF}/{file_name}"
        try:
            req = urllib.request.Request(raw_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as resp: return resp.read()
        except:
            try:
                raw_url_alt = f"https://github.com/{encoded_user}/{encoded_repo}/raw/main/{GITHUB_FOLDER_PDF}/{file_name}"
                req = urllib.request.Request(raw_url_alt, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as resp: return resp.read()
            except: return None

    uploaded_pdfs = st.file_uploader("從本機電腦上傳新講義 PDF (可多選)", type=["pdf"], accept_multiple_files=True)
    total_pdf_count = len(uploaded_pdfs) + len(selected_cloud_pdfs)

    if total_pdf_count > 0:
        st.markdown(f"📊 **目前已鎖定講義總數：{total_pdf_count} 份**")
        
        col_q1, col_q2, col_q3 = st.columns(3)
        with col_q1: page_range = st.text_input("想根據哪幾頁出題？", "整份")
        with col_q2: num_questions = st.number_input("預計生成題數", min_value=1, max_value=100, value=10)
        with col_q3: start_q_num = st.number_input("🔢 設定「起始題號」", min_value=1, max_value=999, value=1, step=1, key="mode_a_qnum")

        st.markdown("---")
        st.subheader("🌐 選擇出題語系樣式")
        lang_style = st.radio(
            "請指定 AI 出題時的題目與選項語言：",
            ["1. 中文出題（專有名詞採『英文 (中文)』雙語標記，貼近國考格式）", "2. 全英文出題（題目與選項皆為 Full English，貼近臨床跑台）"],
            index=0,
            horizontal=True
        )

        st.markdown("---")

        def extract_clean_text_from_pdf(pdf_bytes, pdf_name):
            chunk_text = ""
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for i in range(len(doc)):
                page = doc.load_page(i)
                text = page.get_text()
                if text: chunk_text += f"\n\n=== 【{pdf_name}】第 {i+1} 頁 ===\n{text}"
            return chunk_text

        # ==============================================================================
        # 支線：A' 客製化 Prompt 複製器 (整合「回貼 JSON 文字」直接解析排版)
        # ==============================================================================
        if "功能 A'" in sub_function_mode:
            st.success("💡 專屬客製化 PROMPT 封裝完成！請直接點擊右上角按鈕一鍵複製：")
            combined_text_payload = ""
            for pdf_file in uploaded_pdfs: combined_text_payload += extract_clean_text_from_pdf(pdf_file.read(), pdf_file.name)
            for cloud_pdf_name in selected_cloud_pdfs:
                c_bytes = fetch_cloud_pdf_bytes(cloud_pdf_name)
                if c_bytes: combined_text_payload += extract_clean_text_from_pdf(c_bytes, cloud_pdf_name)
            
            if "1. 中文出題" in lang_style:
                lang_prompt_str = "每個物件中的「題目內容」與「選項A」~「選項E」必須主要使用繁體中文撰寫。遇到醫學專有名詞時，請嚴格採取「英文搭配括號中文」的方式呈現（例如：Myocardial Infarction (心肌梗塞)）。"
            else:
                lang_prompt_str = "每個物件中的「題目內容」與「選項A」~「選項E」必須完全使用純英文 (Full English) 撰寫，符合美國醫學執照考試 (USMLE) 專業醫學出題邏輯。"

            history_prompt_str = ""
            if history_titles:
                history_prompt_str = "\n【🚨 歷史考點去重指令】：以下是歷史已出題目的名單，你設計的新題目絕對禁止再次重複測驗以下已考過的生理機制、藥理靶點或鑑別觀念，必須挑選講義中全新的核心知識點命題：\n" + "\n".join([f"- {t}" for t in history_titles[:40]])

            raw_prompt_for_user = f"""你現在是一位資深的醫學與生物科學教授。請根據我為你提供的這份完整講義文字文本，精準鎖定這些講義文字內容中的【{page_range}】，並圍繞核心主題設計出高質感的題庫。

【數量鐵律】：我要求你精準輸出「剛好」 {num_questions} 題五選一的單選題。絕對不能多出，也不能少出！

【語系要求】：
{lang_prompt_str}

【詳解與出處恆定要求】：
- 不論前面題目是中文還是英文，【針對各選項之詳解】必須一律使用繁體中文進行極為詳細的專家級辨析（逐行解釋為什麼該選項正確或錯誤，換行請用 \\n 符號）。
- 【正確答案】請固定輸出大寫字母（A, B, C, D 或 E）。
- 【出處】請精準對照我文本中的頁碼標籤（例如：=== 【檔名】第 X 頁 ===），指出這題是出自哪一個檔案的第幾頁！
- 正確答案（A, B, C, D, E）的總體數量分布要稍微平均一些。
{history_prompt_str}

【🚨 格式與語法輸出鐵律 - 違者拒收】：
1. 請「只」輸出符合上述規範的標準 JSON 格式列表陣列格式（即以 [ 開頭，以 ] 結尾）。
2. 絕對、嚴禁、不要包含任何 Markdown 包裝字串！例如：禁止在開頭與結尾夾帶 ```json 或 ```。
3. 絕對不要輸出任何多餘的解釋、前言、後記或提示性文字。你的回答必須是 100% 可被機器直接解析的純 JSON 陣列。
4. 嚴格注意物件內最後一個 Key-Value 欄位與最後一個物件的末尾，【絕對不能】有多餘的逗號 (Trailing Comma)。
5. 詳解內容中若需要換行，請務必使用標準字元安全轉義序列「\\\\n」呈現，確保 JSON 的連續性。

【輸出格式規範】：
格式必須是標準的 JSON 格式列表(Array)，內含多個物件，每個物件的 Key 必須嚴格為："題目內容", "選項A", "選項B", "選項C", "選項D", "選項E", "正確答案", "針對各選項之詳解", "出處"

以下是為你夾帶的講義完整純文字文本：
{combined_text_payload}
"""
            st.code(raw_prompt_for_user, language="text")
            
            # 🌟 [全新亮點]：同版面 JSON 回貼排版解析器，讓組員不用跳去模組 C 也能一體化操作
            st.markdown("---")
            st.subheader("📥 🚀 A' 快速回貼解析與雙軸排版引擎")
            st.caption("複製上方 Prompt 至 ChatGPT 或 Claude 等外部 AI。將 AI 產出的 JSON 格式題庫文字，直接貼在下方，即可於此版面一鍵秒速生成 Word 試卷與 Excel 題庫！")
            
            text_input_a_prime = st.text_area("請在下方貼上 AI 吐出的 JSON 陣列內容：", height=250, placeholder='[\n  {\n    "題目內容": "...", \n    "選項A": "..."\n  }\n]', key="a_prime_json_input")
            
            if text_input_a_prime.strip():
                try:
                    clean_text_ap = text_input_a_prime.strip()
                    if clean_text_ap.startswith(BT_JSON): 
                        clean_text_ap = clean_text_ap.split(BT_JSON)[1].split(BT_ONLY)[0].strip()
                    elif clean_text_ap.startswith(BT_ONLY): 
                        clean_text_ap = clean_text_ap.split(BT_ONLY)[1].split(BT_ONLY)[0].strip()
                    
                    json_data_ap = json.loads(clean_text_ap)
                    if isinstance(json_data_ap, list):
                        st.success(f"📝 成功辨識複製貼上的 JSON 文字！偵測到 **{len(json_data_ap)}** 道題目。")
                        
                        # 在 A' 內建專屬的檔名與標題設定欄位
                        st.markdown("---")
                        st.subheader("🏷️ 設定大標題與檔名")
                        end_q_num_ap = start_q_num + len(json_data_ap) - 1
                        calculated_remarks_ap = f"{start_q_num:02d}~{end_q_num_ap:02d}"
                        
                        col_ap1, col_ap2 = st.columns(2)
                        with col_ap1: subject_name_ap = st.text_input("科目名稱", "生理學", key="sub_ap")
                        with col_ap2: teacher_name_ap = st.text_input("老師名稱", "王大明", key="tea_ap")
                            
                        col_ap3, col_ap4 = st.columns(2)
                        with col_ap3: topic_name_ap = st.text_input("課堂主題", "心血管系統", key="top_ap")
                        with col_ap4: remarks_ap = st.text_input("備註 (預設為題號範圍)", value=calculated_remarks_ap, key="rem_ap")
                        
                        final_title_filename_ap = f"{subject_name_ap}_{teacher_name_ap}_{topic_name_ap}_{remarks_ap}"
                        st.info(f"📁 系統預覽輸出名稱將為：**{final_title_filename_ap}**")
                        
                        if st.button("📥 一鍵排版產出 Word 試卷與 Excel 題庫 📥", key="a_prime_convert_btn", use_container_width=True):
                            try:
                                with st.spinner("🎨 正在啟動雙軸排版引擎，同時美化 Word 與 Excel 中..."):
                                    # 1. 產生高質感 Word
                                    doc_ap = Document()
                                    sec_ap = doc_ap.sections[0]
                                    sec_ap.top_margin = sec_ap.bottom_margin = sec_ap.left_margin = sec_ap.right_margin = Cm(1.27)
                                    doc_ap.styles['Normal'].font.name = 'Times New Roman'
                                    doc_ap.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                                    doc_ap.styles['Normal'].font.size = Pt(12)
                                    PURPLE, BLUE = RGBColor(112, 48, 160), RGBColor(0, 50, 150)
                                    
                                    title_p = doc_ap.add_paragraph()
                                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    title_p.add_run(final_title_filename_ap).bold = True
                                    title_p.runs[-1].font.size = Pt(16)
                                    
                                    opt_labels = ['A', 'B', 'C', 'D', 'E']
                                    processed_rows_ap = []
                                    
                                    for idx, item in enumerate(json_data_ap):
                                        current_q_num = int(start_q_num) + idx
                                        
                                        q_txt = item.get("題目內容", item.get("題庫內容", ""))
                                        doc_ap.add_paragraph(f"{current_q_num}. {q_txt}").paragraph_format.space_after = Pt(6)
                                        
                                        for lbl in opt_labels:
                                            opt_txt = item.get(f"選項{lbl}", "")
                                            if opt_txt:
                                                op = doc_ap.add_paragraph(f"({lbl}) {opt_txt}")
                                                op.paragraph_format.left_indent, op.paragraph_format.space_after = Pt(18), Pt(0)
                                                
                                        ans_txt = str(item.get("正確答案", "")).upper().strip()
                                        if ans_txt:
                                            ans_p = doc_ap.add_paragraph()
                                            ans_p.paragraph_format.space_before = Pt(6)
                                            ans_p.add_run("Ans : ").bold = True
                                            ans_p.add_run(f"({ans_txt})")
                                            
                                        expl_txt = str(item.get("針對各選項之詳解", item.get("詳解", ""))).strip()
                                        if expl_txt and expl_txt.lower() != "nan":
                                            h = doc_ap.add_paragraph()
                                            h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(4), Pt(0)
                                            run = h.add_run("詳解 :"); run.bold, run.font.color.rgb = True, PURPLE
                                            
                                            for line in expl_txt.split('\n'):
                                                if not line.strip(): continue
                                                lp = doc_ap.add_paragraph()
                                                lp.paragraph_format.left_indent, lp.paragraph_format.space_after = Pt(18), Pt(2)
                                                m = re.match(r'^([A-F])\s*([\(（].*?[\)隱]|[:：])', line.strip())
                                                if m:
                                                    lp.add_run(m.group(0)).bold = True
                                                    lp.runs[-1].font.color.rgb = PURPLE
                                                    lp.add_run(line.strip()[len(m.group(0)):]).font.color.rgb = PURPLE
                                                else: lp.add_run(line.strip()).font.color.rgb = PURPLE
                                                
                                        src_txt = str(item.get("出處", "")).strip()
                                        if src_txt and src_txt.lower() != "nan":
                                            sp = doc_ap.add_paragraph()
                                            sp.paragraph_format.space_before = Pt(2)
                                            sp.add_run("出處 : ").bold = True
                                            sp.runs[-1].font.color.rgb = BLUE
                                            sp.add_run(src_txt).font.color.rgb = BLUE
                                            
                                        doc_ap.add_paragraph("")
                                        
                                        row_ap = {
                                            '題號': current_q_num,
                                            '題目內容': str(q_txt).strip(),
                                            '選項A': str(item.get("選項A", "")).strip(),
                                            '選項B': str(item.get("選項B", "")).strip(),
                                            '選項C': str(item.get("選項C", "")).strip(),
                                            '選項D': str(item.get("選項D", "")).strip(),
                                            '選項E': str(item.get("選項E", "")).strip(),
                                            '正確答案': ans_txt,
                                            '針對各選項之詳解': expl_txt,
                                            '出處': src_txt
                                        }
                                        processed_rows_ap.append(row_ap)
                                        
                                    # 2. 產生高質感 Excel
                                    excel_out_ap = io.BytesIO()
                                    pd.DataFrame(processed_rows_ap).to_excel(excel_out_ap, index=False)
                                    excel_out_ap.seek(0)
                                    wb_ap = load_workbook(excel_out_ap)
                                    ws_ap = wb_ap.active
                                    for letter, width in EXCEL_COL_WIDTHS.items(): 
                                        ws_ap.column_dimensions[letter].width = width
                                    for r_idx, row in enumerate(ws_ap.iter_rows(min_row=1, max_row=ws_ap.max_row), 1):
                                        for cell in row:
                                            cell.border = EXCEL_BORDER
                                            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                                            if r_idx == 1:
                                                cell.font = Font(bold=True)
                                                cell.alignment = Alignment(horizontal='center', vertical='center')
                                            if cell.column_letter in ['A', 'H']: 
                                                cell.alignment = Alignment(horizontal='center', vertical='center')
                                            if r_idx > 1:
                                                cw = EXCEL_COL_WIDTHS.get(cell.column_letter, 20)
                                                est = math.ceil((len(str(cell.value or '')) * 1.8) / cw)
                                                if est > 1: 
                                                    ws_ap.row_dimensions[r_idx].height = est * 18
                                                    
                                    final_word_bytes_ap = io.BytesIO()
                                    doc_ap.save(final_word_bytes_ap)
                                    
                                    final_excel_bytes_ap = io.BytesIO()
                                    wb_ap.save(final_excel_bytes_ap)
                                    
                                    st.session_state["sol_word_ap"] = final_word_bytes_ap.getvalue()
                                    st.session_state["sol_excel_ap"] = final_excel_bytes_ap.getvalue()
                                    st.session_state["saved_exam_title_ap"] = final_title_filename_ap
                                    
                                    if github_token:
                                        with st.spinner("☁️ 正在即時同步備份 Excel 至 GitHub 歷史庫..."):
                                            success, msg = upload_excel_to_github(final_excel_bytes_ap.getvalue(), f"{final_title_filename_ap}.xlsx", github_token)
                                            if success:
                                                st.success(f"☁️ {msg}")
                                            else:
                                                st.warning(f"⚠️ {msg} (但本地檔案已成功生成)")
                            except Exception as e:
                                st.error(f"轉換排版過程發生錯誤：{e}")
                        
                        # 平行外部下載按鈕
                        if "sol_word_ap" in st.session_state and "sol_excel_ap" in st.session_state:
                            st.success("🎉 Word 試卷與 Excel 題庫排版渲染已完美達成！請點擊下方按鈕下載：")
                            s_name_ap = sanitize_f(st.session_state["saved_exam_title_ap"])
                            dl_col1_ap, dl_col2_ap = st.columns(2)
                            with dl_col1_ap:
                                st.download_button(
                                    label="📊 下載精修 Excel 題庫 (.xlsx)",
                                    data=st.session_state["sol_excel_ap"],
                                    file_name=f"{s_name_ap}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True
                                )
                            with dl_col2_ap:
                                st.download_button(
                                    label="📄 下載精修排版 Word 試卷 (.docx)",
                                    data=st.session_state["sol_word_ap"],
                                    file_name=f"{s_name_ap}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                    else: st.error("❌ 貼上的內容格式不合規，外層必須是標準的方括號列表陣列 `[...]`。")
                except Exception as e: st.error(f"文字 JSON 格式解析失敗，請檢查括號是否完整。錯誤原因: {e}")

        # ==============================================================================
        # 主功能支線：自動在網頁中呼叫 API 跑完 (此功能保留檔名與標題組裝)
        # ==============================================================================
        else:
            st.subheader("🏷️ 設定大標題與檔名")
            end_q_num = start_q_num + num_questions - 1
            default_remarks = f"{start_q_num:02d}~{end_q_num:02d}"

            col_t1, col_t2 = st.columns(2)
            with col_t1: subject_name = st.text_input("科目名稱", "生理學", key="sub_a")
            with col_t2: teacher_name = st.text_input("老師名稱", "王大明", key="tea_a")
                
            col_t3, col_t4 = st.columns(2)
            with col_t3: topic_name = st.text_input("課堂主題", "心血管系統", key="top_a")
            with col_t4: remarks = st.text_input("備註 (預設為題號範圍)", value=default_remarks, key="rem_a")

            final_title_filename = f"{subject_name}_{teacher_name}_{topic_name}_{remarks}"
            st.info(f"📁 系統預覽輸出名稱將為：**{final_title_filename}**")

            if st.button("⚡ 開始全全自動雙模融合出題 ⚡", use_container_width=True):
                try:
                    combined_text_payload = ""
                    with st.spinner("🔍 正在啟動本地高效文字萃取引擎..."):
                        for pdf_file in uploaded_pdfs: combined_text_payload += extract_clean_text_from_pdf(pdf_file.read(), pdf_file.name)
                        for cloud_pdf_name in selected_cloud_pdfs:
                            c_bytes = fetch_cloud_pdf_bytes(cloud_pdf_name)
                            if c_bytes: combined_text_payload += extract_clean_text_from_pdf(c_bytes, cloud_pdf_name)

                    with st.spinner("🧠 任務封裝完成！正透過智慧負載分流引擎發送至雲端核心..."):
                        range_instruction = f"精準鎖定這些講義文字內容中的【{page_range}】" if "整份" not in page_range and "全部" not in page_range else "「通盤掃描並融合這幾份講義文字」的完整內容，宏觀地在不同的章節與核心觀念中提取重點"
                        history_block = ""
                        if history_titles: 
                            history_block = """
                            【🚨 歷史考點去重(Tag-based De-duplication) 嚴格指令】：
                            以下是我提供給你的歷史已出題庫列表。你身為資深教授，請深入分析以下每一道舊題目背後的『核心醫學考點、生理機制、臨床症狀、藥理靶點或診斷核心』：
                            """ + "\n".join([f"- 舊題範例: {t}" for t in history_titles[:80]]) + """
                            【核心鐵律】：
                            1. 分析完上述舊題目的核心考點後，本次設計的新題目『絕對禁止』再次重複測驗這些已考過的機制！
                            2. 請從下方的講義文本中，發掘全新、尚未被上述歷史題目覆蓋的生理機轉、鑑別診斷或臨床指標來進行命題。
                            """

                        if "1. 中文出題" in lang_style:
                            lang_prompt = "每個物件中的「題目內容」與「選項A」~「選項E」必須主要使用繁體中文撰寫。遇到醫學專有名詞時，請嚴格採取「英文搭配括號中文」的方式呈現（例如：Myocardial Infarction (心肌梗塞)）。"
                        else:
                            lang_prompt = "每個物件中的「題目內容」與「選項A」~「選項E」必須完全使用純英文 (Full English) 寫作。符合美國醫學執照考試 (USMLE) 專業醫學出題邏輯。"

                        prompt = f"""你現在是一位資深的醫學與生物科學教授。請根據我為你提供的這份【完整】講義文字內容，{range_instruction}，並圍繞核心主題【{topic_name}】出題。
                        【數量鐵律】：我要求你精準輸出「剛好」 {num_questions} 題五選一的單選題。絕對不能多出，也不能少出！
                        {lang_prompt}
                        {history_block}
                        【詳解與出處恆定要求】：
                        - 不論前面題目是中文還是英文，【針對各選項之詳解】必須一律使用繁體中文進行極為詳細的專家級辨析（逐行解釋為什麼該選項正確或錯誤，換行請用 \\n 符號）。
                        - 【正確答案】請固定輸出大寫字母（A, B, C, D 或 E）。
                        - 【出處】請精準對照我文本中的頁碼標籤（例如：=== 【檔名】第 X 頁 ===），指出這題是出自哪一個檔案的第幾頁！
                        - 正確答案（A, B, C, D, E）的總體數量分布要稍微平均一些。
                        
                        【🚨 格式與語法輸出鐵律 - 違者拒收】：
                        1. 請「只」輸出標準 JSON 格式列表陣列格式（即以 [ 開頭，以 ] 結尾）。
                        2. 絕對、嚴禁、不要包含 any Markdown 外包裝字串！禁止在開頭與結尾夾帶 ```json 等字眼。
                        3. 絕對不要輸出 any 多餘的解釋、前言或後記。
                        4. 嚴格注意物件內最後一個欄位與最後一個物件的末尾，【絕對不能】有多餘的逗號。
                        5. 詳解換行請務必使用安全轉義序列「\\\\n」呈現。

                        格式必須是 JSON 格式的列表(Array)，內含多個物件，每個物件的Key必須嚴格為："題目內容", "選項A", "選項B", "選項C", "選項D", "選項E", "正確答案", "針對各選項之詳解", "出處"

                        以下是待讀取的講義完整純文字文本：
                        {combined_text_payload}
                        """

                        clean_response = generate_content_via_http_with_retry([prompt], api_key)
                        clean_response = clean_response.strip()
                        if clean_response.startswith(BT_JSON): clean_response = clean_response.split(BT_JSON)[1].split(BT_ONLY)[0].strip()
                        elif clean_response.startswith(BT_ONLY): clean_response = clean_response.split(BT_ONLY)[1].split(BT_ONLY)[0].strip()
                            
                        raw_questions = json.loads(clean_response)
                        raw_questions = raw_questions[:num_questions]

                    with st.spinner("🎨 題目設計完成！正在套用高質感格式排版引擎..."):
                        processed_rows = []
                        opt_labels = ['A', 'B', 'C', 'D', 'E']
                        for idx, q in enumerate(raw_questions):
                            current_q_num = int(start_q_num) + idx
                            row_dict = {'題號': current_q_num, '題目內容': str(q.get('題目內容', '')).strip()}
                            for lbl in opt_labels: row_dict[f'選項{lbl}'] = str(q.get(f'選項{lbl}', '')).strip()
                            ans = str(q.get('正確答案', '')).upper().strip()
                            row_dict['正確答案'] = ans if ans in opt_labels else ""
                            row_dict['針對各選項之詳解'] = str(q.get('針對各選項之詳解', '')).strip()
                            row_dict['出處'] = str(q.get('出處', '')).strip()
                            processed_rows.append(row_dict)

                        excel_out = io.BytesIO()
                        pd.DataFrame(processed_rows).to_excel(excel_out, index=False)
                        excel_out.seek(0)
                        wb = load_workbook(excel_out)
                        ws = wb.active
                        for letter, width in EXCEL_COL_WIDTHS.items(): ws.column_dimensions[letter].width = width
                        for r_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row), 1):
                            for cell in row:
                                cell.border = EXCEL_BORDER
                                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                                if r_idx == 1:
                                    cell.font = Font(bold=True)
                                    cell.alignment = Alignment(horizontal='center', vertical='center')
                                if cell.column_letter in ['A', 'H']: cell.alignment = Alignment(horizontal='center', vertical='center')
                                if r_idx > 1:
                                    cw = EXCEL_COL_WIDTHS.get(cell.column_letter, 20)
                                    est = math.ceil((len(str(cell.value or '')) * 1.8) / cw)
                                    if est > 1: ws.row_dimensions[r_idx].height = est * 18
                        final_excel_bytes = io.BytesIO()
                        wb.save(final_excel_bytes)

                        doc = Document()
                        sec = doc.sections[0]
                        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1.27)
                        doc.styles['Normal'].font.name = 'Times New Roman'
                        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                        doc.styles['Normal'].font.size = Pt(12)
                        PURPLE, BLUE = RGBColor(112, 48, 160), RGBColor(0, 50, 150)
                        
                        title_p = doc.add_paragraph()
                        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        title_p.add_run(final_title_filename).bold = True
                        title_p.runs[-1].font.size = Pt(16)

                        for r in processed_rows:
                            doc.add_paragraph(f"{r['題號']}. {r['題目內容']}").paragraph_format.space_after = Pt(6)
                            for lbl in opt_labels:
                                txt = r.get(f'選項{lbl}', '')
                                if txt:
                                    op = doc.add_paragraph(f"({lbl}) {txt}")
                                    op.paragraph_format.left_indent, op.paragraph_format.space_after = Pt(18), Pt(0)
                            ans_p = doc.add_paragraph()
                            ans_p.paragraph_format.space_before = Pt(6)
                            ans_p.add_run("Ans : ").bold = True
                            ans_p.add_run(f"({r['正確答案']})")
                            expl = str(r['針對各選項之詳解'])
                            if expl and expl.lower() != "nan":
                                h = doc.add_paragraph()
                                h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(4), Pt(0)
                                run = h.add_run("詳解 :"); run.bold, run.font.color.rgb = True, PURPLE
                                for line in expl.split('\n'):
                                    if not line.strip(): continue
                                    lp = doc.add_paragraph()
                                    lp.paragraph_format.left_indent, lp.paragraph_format.space_after = Pt(18), Pt(2)
                                    m = re.match(r'^([A-F])\s*([\(（].*?[\)隱]|[:：])', line.strip())
                                    if m:
                                        lp.add_run(m.group(0)).bold = True
                                        lp.runs[-1].font.color.rgb = PURPLE
                                        lp.add_run(line.strip()[len(m.group(0)):]).font.color.rgb = PURPLE
                                    else: lp.add_run(line.strip()).font.color.rgb = PURPLE
                            src = str(r['出處'])
                            if src and src.lower() != "nan":
                                sp = doc.add_paragraph()
                                sp.paragraph_format.space_before = Pt(2)
                                sp.add_run("出處 : ").bold = True
                                sp.runs[-1].font.color.rgb = BLUE
                                sp.add_run(src).font.color.rgb = BLUE
                            doc.add_paragraph("")

                        final_word_bytes = io.BytesIO()
                        doc.save(final_word_bytes)
                        st.session_state["generated_excel_a"] = final_excel_bytes.getvalue()
                        st.session_state["generated_word_a"] = final_word_bytes.getvalue()
                        st.session_state["saved_exam_title_a"] = final_title_filename
                        
                        # 🌟 [即時同步]：若設定了 Token，自動推送至 GitHub 歷史庫
                        if github_token:
                            with st.spinner("☁️ 正在即時同步備份 Excel 至 GitHub 歷史庫..."):
                                success, msg = upload_excel_to_github(final_excel_bytes.getvalue(), f"{final_title_filename}.xlsx", github_token)
                                if success:
                                    st.success(f"☁️ {msg}")
                                else:
                                    st.warning(f"⚠️ {msg} (但本地檔案已成功生成)")
                                    
                except Exception as e: st.error(f"出題過程出錯：{e}")

            if "generated_excel_a" in st.session_state and "generated_word_a" in st.session_state:
                st.success("🎉 模式 A：講義題庫與試卷皆已設計完成！請下載：")
                s_name = sanitize_f(st.session_state["saved_exam_title_a"])
                dl_col1, dl_col2 = st.columns(2)
                with dl_col1: st.download_button("📊 下載精修 Excel 題庫 (.xlsx)", data=st.session_state["generated_excel_a"], file_name=f"{s_name}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                with dl_col2: st.download_button("📄 下載精修 Word 試卷 (.docx)", data=st.session_state["generated_word_a"], file_name=f"{s_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

# ==============================================================================
# 🌟 模組 B：現成題目自動配詳解系統
# ==============================================================================
elif "模組 B" in main_mode:
    st.subheader("📝 模式 B：現成題目自動配詳解")
    uploaded_file = st.file_uploader("請選擇包含純題目的 Excel 檔案 (.xlsx)", type=["xlsx"], key="mode_b_uploader")
    if uploaded_file:
        try:
            df_input = pd.read_excel(uploaded_file)
            st.success(f"📊 成功讀取現有考題！共偵測到 **{len(df_input)}** 道題目。")
            
            col_q = next((c for c in df_input.columns if any(k in str(c) for k in ["題目", "Question", "內容"])), None)
            col_a = next((c for c in df_input.columns if "A" in str(c)), None)
            col_b = next((c for c in df_input.columns if "B" in str(c)), None)
            col_c = next((c for c in df_input.columns if "C" in str(c)), None)
            col_d = next((c for c in df_input.columns if "D" in str(c)), None)
            col_e = next((c for c in df_input.columns if "E" in str(c)), None)

            if not (col_q and col_a and col_b):
                st.error("❌ Excel 內找不到對應的『題目內容』或『選項』表頭欄位。")
                st.stop()

            start_q_num_b = st.number_input("🔢 設定「起始題號」", min_value=1, max_value=999, value=1, step=1, key="mode_b_qnum")
            end_q_num_b = start_q_num_b + len(df_input) - 1
            calculated_remarks_b = f"{start_q_num_b:02d}~{end_q_num_b:02d}"

            st.markdown("---")
            st.subheader("🏷️ 設定大標題與檔名")
            col_t1_b, col_t2_b = st.columns(2)
            with col_t1_b: subject_name_b = st.text_input("科目名稱", "生理學", key="sub_b")
            with col_t2_b: teacher_name_b = st.text_input("老師名稱", "王大明", key="tea_b")
            col_t3_b, col_t4_b = st.columns(2)
            with col_t3_b: topic_name_b = st.text_input("課堂主題", "心血管系統", key="top_b")
            with col_t4_b: remarks_b = st.text_input("備註 (預設為題號範圍)", value=calculated_remarks_b, key="rem_b")

            final_title_filename_b = f"{subject_name_b}_{teacher_name_b}_{topic_name_b}_{remarks_b}"
            st.info(f"📁 系統預覽輸出名稱將為：**{final_title_filename_b}**")

            cleaned_questions = []
            for idx, row in df_input.iterrows():
                cleaned_questions.append({
                    "orig_idx": idx + 1,
                    "題目內容": str(row[col_q]).strip(),
                    "選項A": str(row[col_a]).strip() if pd.notna(row[col_a]) else "",
                    "選項B": str(row[col_b]).strip() if pd.notna(row[col_b]) else "",
                    "選項C": str(row[col_c]).strip() if pd.notna(row[col_c]) else "",
                    "選項D": str(row[col_d]).strip() if pd.notna(row[col_d]) else "",
                    "選項E": str(row[col_e]).strip() if col_e and pd.notna(row[col_e]) else ""
                })

            if st.button("⚡ 開始全自動配對醫學詳解 ⚡", use_container_width=True):
                try:
                    with st.spinner("🧠 任務封裝完成！正在跨世代智慧調度配對詳解中..."):
                        input_data_json = json.dumps(cleaned_questions, ensure_ascii=False)
                        prompt = f"""你現在是一位資深的醫學與生物科學教授。請根據我提供給你的 JSON 題目列表，【原封不動】地保留題目內容與選項，並補上最精準的【正確答案】以及極為詳細的【針對各選項之詳解】。
                        
                        【🚨 格式與語法輸出鐵律 - 違者拒收】：
                        1. 請「只」輸出標準 JSON 格式列表陣列格式（即以 [ 開頭，以 ] 結尾）。
                        2. 絕對、嚴禁、不要包含 any Markdown 外包裝字串！禁止在開頭與結尾夾帶 ```json 等字眼。
                        3. 絕對不要輸出任何多餘的解釋、前言或後記。
                        4. 嚴格注意物件內最後一個欄位與最後一個物件的末尾，【絕對不能】有多餘的逗號。
                        5. 詳解換行請務必使用安全轉義序列「\\\\n」呈現。

                        格式必須嚴格符合 JSON 列表(Array)，Key 必須為："題目內容", "選項A", "選項B", "選項C", "選項D", "選項E", "正確答案", "針對各選項之詳解"
                        {input_data_json}"""
                        
                        ai_response = generate_content_via_http_with_retry([prompt], api_key)
                        ai_response = ai_response.strip()
                        if ai_response.startswith(BT_JSON): ai_response = ai_response.split(BT_JSON)[1].split(BT_ONLY)[0].strip()
                        elif ai_response.startswith(BT_ONLY): ai_response = ai_response.split(BT_ONLY)[1].split(BT_ONLY)[0].strip()
                        raw_results = json.loads(ai_response)

                    with st.spinner("🎨 詳解補全完成！正在套用高質感格式排版引擎..."):
                        processed_rows_b = []
                        opt_labels = ['A', 'B', 'C', 'D', 'E']
                        for idx, q in enumerate(raw_results):
                            current_q_num = int(start_q_num_b) + idx
                            row_dict = {'題號': current_q_num, '題目內容': str(q.get('題目內容', '')).strip()}
                            for lbl in opt_labels: row_dict[f'選項{lbl}'] = str(q.get(f'選項{lbl}', '')).strip()
                            ans = str(q.get('正確答案', '')).upper().strip()
                            row_dict['正確答案'] = ans if ans in opt_labels else ""
                            row_dict['針對各選項之詳解'] = str(q.get('針對各選項之詳解', '')).strip()
                            row_dict['出處'] = "醫學資料庫專家詳解"
                            processed_rows_b.append(row_dict)

                        excel_out_b = io.BytesIO()
                        pd.DataFrame(processed_rows_b).to_excel(excel_out_b, index=False)
                        excel_out_b.seek(0)
                        wb_b = load_workbook(excel_out_b)
                        ws_b = wb_b.active
                        for letter, width in EXCEL_COL_WIDTHS.items(): ws_b.column_dimensions[letter].width = width
                        for r_idx, row in enumerate(ws_b.iter_rows(min_row=1, max_row=ws_b.max_row), 1):
                            for cell in row:
                                cell.border = EXCEL_BORDER
                                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                                if r_idx == 1:
                                    cell.font = Font(bold=True)
                                    cell.alignment = Alignment(horizontal='center', vertical='center')
                                if cell.column_letter in ['A', 'H']: cell.alignment = Alignment(horizontal='center', vertical='center')
                                if r_idx > 1:
                                    cw = EXCEL_COL_WIDTHS.get(cell.column_letter, 20)
                                    est = math.ceil((len(str(cell.value or '')) * 1.8) / cw)
                                    if est > 1: ws_b.row_dimensions[r_idx].height = est * 18
                        final_excel_bytes_b = io.BytesIO()
                        wb_b.save(final_excel_bytes_b)

                        doc_b = Document()
                        sec_b = doc_b.sections[0]
                        sec_b.top_margin = sec_b.bottom_margin = sec_b.left_margin = sec_b.right_margin = Cm(1.27)
                        doc_b.styles['Normal'].font.name = 'Times New Roman'
                        doc_b.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                        doc_b.styles['Normal'].font.size = Pt(12)
                        PURPLE = RGBColor(112, 48, 160)
                        
                        title_p_b = doc_b.add_paragraph()
                        title_p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        title_p_b.add_run(final_title_filename_b).bold = True
                        title_p_b.runs[-1].font.size = Pt(16)

                        for r in processed_rows_b:
                            doc_b.add_paragraph(f"{r['題號']}. {r['題目內容']}").paragraph_format.space_after = Pt(6)
                            for lbl in opt_labels:
                                txt = r.get(f'選項{lbl}', '')
                                if txt:
                                    op = doc_b.add_paragraph(f"({lbl}) {txt}")
                                    op.paragraph_format.left_indent, op.paragraph_format.space_after = Pt(18), Pt(0)
                            ans_p = doc_b.add_paragraph()
                            ans_p.paragraph_format.space_before = Pt(6)
                            ans_p.add_run("Ans : ").bold = True
                            ans_p.add_run(f"({r['正確答案']})")
                            expl = str(r['針對各選項之詳解'])
                            if expl and expl.lower() != "nan":
                                h = doc_b.add_paragraph()
                                h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(4), Pt(0)
                                run = h.add_run("詳解 :"); run.bold, run.font.color.rgb = True, PURPLE
                                for line in expl.split('\n'):
                                    if not line.strip(): continue
                                    lp = doc_b.add_paragraph()
                                    lp.paragraph_format.left_indent, lp.paragraph_format.space_after = Pt(18), Pt(2)
                                    m = re.match(r'^([A-F])\s*([\(（].*?[\)隱]|[:：])', line.strip())
                                    if m:
                                        lp.add_run(m.group(0)).bold = True
                                        lp.runs[-1].font.color.rgb = PURPLE
                                        lp.add_run(line.strip()[len(m.group(0)):]).font.color.rgb = PURPLE
                                    else: lp.add_run(line.strip()).font.color.rgb = PURPLE
                            doc_b.add_paragraph("")

                        final_word_bytes_b = io.BytesIO()
                        doc_b.save(final_word_bytes_b)
                        st.session_state["sol_excel_b"] = final_excel_bytes_b.getvalue()
                        st.session_state["sol_word_b"] = final_word_bytes_b.getvalue()
                        st.session_state["saved_exam_title_b"] = final_title_filename_b
                        
                        # 🌟 [即時同步]：若設定了 Token，自動推送至 GitHub 歷史庫
                        if github_token:
                            with st.spinner("☁️ 正在即時同步備份 Excel 至 GitHub 歷史庫..."):
                                success, msg = upload_excel_to_github(final_excel_bytes_b.getvalue(), f"{final_title_filename_b}.xlsx", github_token)
                                if success:
                                    st.success(f"☁️ {msg}")
                                else:
                                    st.warning(f"⚠️ {msg} (但本地檔案已成功生成)")
                                    
                except Exception as e: st.error(f"分析過程出錯：{e}")

            if "sol_excel_b" in st.session_state and "sol_word_b" in st.session_state:
                st.success("🎉 模式 B 處理完畢！請下載：")
                s_name_b = sanitize_f(st.session_state["saved_exam_title_b"])
                dl_col1_b, dl_col2_b = st.columns(2)
                with dl_col1_b: st.download_button("📊 下載附詳解題庫 (.xlsx)", data=st.session_state["sol_excel_b"], file_name=f"{s_name_b}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                with dl_col2_b: st.download_button("📄 下載附詳解試卷 (.docx)", data=st.session_state["sol_word_b"], file_name=f"{s_name_b}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        except Exception as e: st.error(f"讀取 Excel 檔案錯誤：{e}")

# ==============================================================================
# 🌟 模組 C：既有題庫已含詳解多管道輸入系統 (Excel / JSON 檔案)
# ==============================================================================
else:
    st.subheader("📄 模式 C：既有題庫 ➡️ 高品質渲染 Word/Excel 考卷 (免金鑰)")
    
    input_channel = st.radio(
        "📥 請選擇您的資料輸入來源：",
        ["1. 上傳既有含詳解的 Excel (.xlsx)", "2. 上傳 AI 吐出的 JSON 檔案 (.json)"],
        index=0,
        horizontal=True
    )
    st.markdown("---")

    raw_items_list = [] 

    if "1. 上傳既有" in input_channel:
        uploaded_file_c = st.file_uploader("請選擇 Excel 檔案 (.xlsx)", type=["xlsx"], key="c_xlsx")
        if uploaded_file_c:
            try:
                df_c = pd.read_excel(uploaded_file_c)
                col_q = next((c for c in df_c.columns if any(k in str(c) for k in ["題目", "Question", "內容"])), None)
                col_a = next((c for c in df_c.columns if "A" in str(c)), None)
                col_b = next((c for c in df_c.columns if "B" in str(c)), None)
                col_c = next((c for c in df_c.columns if "C" in str(c)), None)
                col_d = next((c for c in df_c.columns if "D" in str(c)), None)
                col_e = next((c for c in df_c.columns if "E" in str(c)), None)
                col_ans = next((c for c in df_c.columns if any(k in str(c) for k in ["答案", "正確", "Answer"])), None)
                col_expl = next((c for c in df_c.columns if any(k in str(c) for k in ["詳解", "解析", "Explain", "Explanation"])), None)
                col_src = next((c for c in df_c.columns if any(k in str(c) for k in ["出處", "來源", "Source"])), None)
                
                if col_q and col_a and col_b:
                    for _, row in df_c.iterrows():
                        raw_items_list.append({
                            "題目內容": str(row[col_q]).strip(),
                            "選項A": str(row[col_a]).strip() if pd.notna(row[col_a]) else "",
                            "選項B": str(row[col_b]).strip() if pd.notna(row[col_b]) else "",
                            "選項C": str(row[col_c]).strip() if pd.notna(row[col_c]) else "",
                            "選項D": str(row[col_d]).strip() if pd.notna(row[col_d]) else "",
                            "選項E": str(row[col_e]).strip() if col_e and pd.notna(row[col_e]) else "",
                            "正確答案": str(row[col_ans]).strip() if col_ans and pd.notna(row[col_ans]) else "",
                            "針對各選項之詳解": str(row[col_expl]).strip() if col_expl and pd.notna(row[col_expl]) else "",
                            "出處": str(row[col_src]).strip() if col_src and pd.notna(row[col_src]) else ""
                        })
                    st.success(f"📊 成功加載 Excel 題庫！偵測到 **{len(raw_items_list)}** 道題目。")
            except Exception as e: st.error(f"Excel 解析失敗：{e}")

    else:
        uploaded_json_c = st.file_uploader("請選擇 JSON 檔案 (.json)", type=["json"], key="c_json")
        if uploaded_json_c:
            try:
                json_data = json.load(uploaded_json_c)
                if isinstance(json_data, list):
                    raw_items_list = json_data
                    st.success(f"📂 成功加載 JSON 檔案！偵測到 **{len(raw_items_list)}** 道題目。")
                else: st.error("❌ JSON 格式不正確，外層必須是一個列表 (Array)。")
            except Exception as e: st.error(f"JSON 檔案讀取失敗：{e}")

    if len(raw_items_list) > 0:
        start_q_num_c = st.number_input("🔢 設定「起始題號」", min_value=1, max_value=999, value=1, step=1, key="mode_c_qnum")
        end_q_num_c = start_q_num_c + len(raw_items_list) - 1
        calculated_remarks_c = f"{start_q_num_c:02d}~{end_q_num_c:02d}"

        st.markdown("---")
        st.subheader("🏷️ 設定大標題與檔名")
        col_t1_c, col_t2_c = st.columns(2)
        with col_t1_c: subject_name_c = st.text_input("科目名稱", "生理學", key="sub_c")
        with col_t2_c: teacher_name_c = st.text_input("老師名稱", "王大明", key="tea_c")
            
        col_t3_c, col_t4_c = st.columns(2)
        with col_t3_c: topic_name_c = st.text_input("課堂主題", "心血管系統", key="top_c")
        with col_t4_c: remarks_c = st.text_input("備註 (預設為題號範圍)", value=calculated_remarks_c, key="rem_c")

        final_title_filename_c = f"{subject_name_c}_{teacher_name_c}_{topic_name_c}_{remarks_c}"
        st.info(f"📁 系統預覽輸出名稱將為：**{final_title_filename_c}**")

        if st.button("📥 一鍵排版產出 Word 試卷與 Excel 題庫 📥", use_container_width=True):
            try:
                with st.spinner("🎨 正在啟動雙軸排版引擎，同時美化 Word 與 Excel 中..."):
                    # 1. 產生高質感 Word
                    doc_c = Document()
                    sec_c = doc_c.sections[0]
                    sec_c.top_margin = sec_c.bottom_margin = sec_c.left_margin = sec_c.right_margin = Cm(1.27)
                    doc_c.styles['Normal'].font.name = 'Times New Roman'
                    doc_c.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                    doc_c.styles['Normal'].font.size = Pt(12)
                    
                    PURPLE, BLUE = RGBColor(112, 48, 160), RGBColor(0, 50, 150)
                    
                    title_p = doc_c.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_p.add_run(final_title_filename_c).bold = True
                    title_p.runs[-1].font.size = Pt(16)

                    opt_labels = ['A', 'B', 'C', 'D', 'E']
                    processed_rows_c = []

                    for idx, item in enumerate(raw_items_list):
                        current_q_num = int(start_q_num_c) + idx
                        
                        # --- 寫入 Word ---
                        q_txt = item.get("題目內容", item.get("題庫內容", ""))
                        doc_c.add_paragraph(f"{current_q_num}. {q_txt}").paragraph_format.space_after = Pt(6)
                        
                        for lbl in opt_labels:
                            opt_txt = item.get(f"選項{lbl}", "")
                            if opt_txt:
                                op = doc_c.add_paragraph(f"({lbl}) {opt_txt}")
                                op.paragraph_format.left_indent, op.paragraph_format.space_after = Pt(18), Pt(0)
                        
                        ans_txt = str(item.get("正確答案", "")).upper().strip()
                        if ans_txt:
                            ans_p = doc_c.add_paragraph()
                            ans_p.paragraph_format.space_before = Pt(6)
                            ans_p.add_run("Ans : ").bold = True
                            ans_p.add_run(f"({ans_txt})")
                        
                        expl_txt = str(item.get("針對各選項之詳解", item.get("詳解", ""))).strip()
                        if expl_txt and expl_txt.lower() != "nan":
                            h = doc_c.add_paragraph()
                            h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(4), Pt(0)
                            run = h.add_run("詳解 :"); run.bold, run.font.color.rgb = True, PURPLE
                            
                            for line in expl_txt.split('\n'):
                                if not line.strip(): continue
                                lp = doc_c.add_paragraph()
                                lp.paragraph_format.left_indent, lp.paragraph_format.space_after = Pt(18), Pt(2)
                                m = re.match(r'^([A-F])\s*([\(（].*?[\)隱]|[:：])', line.strip())
                                if m:
                                    lp.add_run(m.group(0)).bold = True
                                    lp.runs[-1].font.color.rgb = PURPLE
                                    lp.add_run(line.strip()[len(m.group(0)):]).font.color.rgb = PURPLE
                                else: lp.add_run(line.strip()).font.color.rgb = PURPLE
                        
                        src_txt = str(item.get("出處", "")).strip()
                        if src_txt and src_txt.lower() != "nan":
                            sp = doc_c.add_paragraph()
                            sp.paragraph_format.space_before = Pt(2)
                            sp.add_run("出處 : ").bold = True
                            sp.runs[-1].font.color.rgb = BLUE
                            sp.add_run(src_txt).font.color.rgb = BLUE
                            
                        doc_c.add_paragraph("")
                        
                        # --- 整理資料準備給 Excel ---
                        row_dict = {
                            '題號': current_q_num,
                            '題目內容': str(q_txt).strip(),
                            '選項A': str(item.get("選項A", "")).strip(),
                            '選項B': str(item.get("選項B", "")).strip(),
                            '選項C': str(item.get("選項C", "")).strip(),
                            '選項D': str(item.get("選項D", "")).strip(),
                            '選項E': str(item.get("選項E", "")).strip(),
                            '正確答案': ans_txt,
                            '針對各選項之詳解': expl_txt,
                            '出處': src_txt
                        }
                        processed_rows_c.append(row_dict)

                    # 2. 產生高質感 Excel
                    excel_out_c = io.BytesIO()
                    pd.DataFrame(processed_rows_c).to_excel(excel_out_c, index=False)
                    excel_out_c.seek(0)
                    wb_c = load_workbook(excel_out_c)
                    ws_c = wb_c.active
                    for letter, width in EXCEL_COL_WIDTHS.items(): 
                        ws_c.column_dimensions[letter].width = width
                    for r_idx, row in enumerate(ws_c.iter_rows(min_row=1, max_row=ws_c.max_row), 1):
                        for cell in row:
                            cell.border = EXCEL_BORDER
                            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                            if r_idx == 1:
                                cell.font = Font(bold=True)
                                cell.alignment = Alignment(horizontal='center', vertical='center')
                            if cell.column_letter in ['A', 'H']: 
                                cell.alignment = Alignment(horizontal='center', vertical='center')
                            if r_idx > 1:
                                cw = EXCEL_COL_WIDTHS.get(cell.column_letter, 20)
                                est = math.ceil((len(str(cell.value or '')) * 1.8) / cw)
                                if est > 1: 
                                    ws_c.row_dimensions[r_idx].height = est * 18

                    final_word_bytes_c = io.BytesIO()
                    doc_c.save(final_word_bytes_c)
                    
                    final_excel_bytes_c = io.BytesIO()
                    wb_c.save(final_excel_bytes_c)
                    
                    # 寫入狀態
                    st.session_state["sol_word_c"] = final_word_bytes_c.getvalue()
                    st.session_state["sol_excel_c"] = final_excel_bytes_c.getvalue()
                    st.session_state["saved_exam_title_c"] = final_title_filename_c
                    
                    # 🌟 [即時同步]：若設定了 Token，自動推送至 GitHub 歷史庫
                    if github_token:
                        with st.spinner("☁️ 正在即時同步備份 Excel 至 GitHub 歷史庫..."):
                            success, msg = upload_excel_to_github(final_excel_bytes_c.getvalue(), f"{final_title_filename_c}.xlsx", github_token)
                            if success:
                                st.success(f"☁️ {msg}")
                            else:
                                st.warning(f"⚠️ {msg} (但本地檔案已成功生成)")
                                
            except Exception as e:
                st.error(f"轉換排版過程發生錯誤：{e}")

        # 下載按鈕 (平級抽離，完全獨立於 try-except 外部，安全且完美對齊)
        if "sol_word_c" in st.session_state and "sol_excel_c" in st.session_state:
            st.success("🎉 Word 試卷與 Excel 題庫排版渲染已完美達成！請點擊下方按鈕下載：")
            s_name_c = sanitize_f(st.session_state["saved_exam_title_c"])
            dl_col1_c, dl_col2_c = st.columns(2)
            with dl_col1_c:
                st.download_button(
                    label="📊 下載精修 Excel 題庫 (.xlsx)",
                    data=st.session_state["sol_excel_c"],
                    file_name=f"{s_name_c}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            with dl_col2_c:
                st.download_button(
                    label="📄 下載精修排版 Word 試卷 (.docx)",
                    data=st.session_state["sol_word_c"],
                    file_name=f"{s_name_c}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
